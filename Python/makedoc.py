# Re-attempt to generate and save the document
from docx import Document

# Create a new Word Document
doc = Document()

# Add Title
doc.add_heading('Answers to Questions on AC Circuit Analysis and Complex Numbers', 0)

# Add Q3.1
doc.add_heading('Q 3.1: How does the concept of impedance relate to the use of complex numbers in AC circuit analysis?', level=1)
doc.add_paragraph(
    "Impedance (Z) in AC circuits is a generalization of resistance that accounts for both resistance (R) "
    "and reactance (X). It is represented using complex numbers:"
)
doc.add_paragraph("Z = R + jX", style='Quote')
doc.add_paragraph(
    "- R: Resistance (real part)\n"
    "- X: Reactance (imaginary part), which could be inductive or capacitive."
)
doc.add_paragraph(
    "Impedance allows engineers to easily represent and calculate the behavior of components such as resistors, "
    "capacitors, and inductors:\n"
    "- Inductor: Z_L = jωL\n"
    "- Capacitor: Z_C = -j / (ωC)"
)
doc.add_paragraph(
    "Using complex numbers makes it possible to:\n"
    "1. Analyze phase relationships between voltage and current.\n"
    "2. Add impedances in series or parallel just like adding resistances in DC circuits.\n"
    "3. Calculate current and voltage in AC circuits using Ohm’s Law in complex form: "
)
doc.add_paragraph("I = V / Z", style='Quote')

# Add Q3.2
doc.add_heading('Q 3.2: How does changing the frequency of the AC source affect the impedance of inductors and capacitors, and how is this reflected in the calculations?', level=1)
doc.add_paragraph(
    "The impedance of inductors and capacitors depends on the frequency of the AC source (ω = 2πf):"
)
doc.add_paragraph(
    "1. Inductor:\n"
    "   Z_L = jωL = j(2πf)L\n"
    "   As the frequency (f) increases, the impedance of the inductor increases.\n\n"
    "2. Capacitor:\n"
    "   Z_C = -j / (ωC) = -j / (2πfC)\n"
    "   As the frequency (f) increases, the impedance of the capacitor decreases."
)
doc.add_paragraph(
    "At low frequencies, inductors behave like short circuits (low impedance) and capacitors behave like open circuits (high impedance). "
    "At high frequencies, inductors behave like open circuits (high impedance) and capacitors behave like short circuits (low impedance)."
)

# Add Q3.3
doc.add_heading('Q 3.3: How would you extend this program to analyze circuits with multiple resistors, capacitors, and inductors in both series and parallel configurations?', level=1)
doc.add_paragraph(
    "To extend the program to handle circuits with multiple resistors, capacitors, and inductors in both series and parallel, you can:"
)
doc.add_paragraph(
    "1. Implement a `ParallelCircuit` class:\n"
    "   - Create a class derived from `SeriesCircuit` that calculates the total impedance for components connected in parallel.\n"
    "   - For parallel components: 1/Z_total = 1/Z1 + 1/Z2 + ... + 1/Zn\n\n"
    "2. Modify the `SeriesCircuit` class to store multiple components in series. The total impedance is simply the sum: Z_total = Z1 + Z2 + ... + Zn\n"
    "3. Add functionality for mixed circuits, allowing the user to nest series and parallel configurations.\n"
    "4. Include methods to recalculate impedance dynamically when the frequency changes."
)

# Add Q3.4
doc.add_heading('Q 3.4: What practical applications does complex number arithmetic have in other areas of electrical engineering?', level=1)
doc.add_paragraph(
    "1. Signal Processing:\n"
    "   - Complex numbers represent sinusoidal signals in the frequency domain (e.g., Fourier transforms).\n"
    "   - Phasor representation simplifies calculations of amplitude and phase.\n\n"
    "2. Control Systems:\n"
    "   - Transfer functions and root locus plots use complex numbers for system analysis.\n\n"
    "3. Communication Systems:\n"
    "   - Complex modulation schemes (e.g., QAM and OFDM) rely on complex arithmetic.\n"
    "   - Fourier analysis decomposes signals for transmission.\n\n"
    "4. Power Systems:\n"
    "   - Power flows are analyzed using complex power: S = P + jQ (where P is real power, and Q is reactive power).\n\n"
    "5. Electromagnetic Field Analysis:\n"
    "   - Maxwell's equations involve complex numbers to represent electromagnetic waves and fields."
)

# Save the document
file_path = "C:/Users/Kanna/Documents/AC_Circuit_Analysis_Answers_Final.docx"
doc.save(file_path)

file_path